"""Analyze docx files to understand their structure."""
from docx import Document
from docx.oxml.ns import qn
import os

FILES = [
    r"C:\Users\王文超\Desktop\Servre_SIT_Mastery\docs\100-day-plan\day-docs\Day1.docx",
    r"C:\Users\王文超\Desktop\Servre_SIT_Mastery\docs\100-day-plan\day-docs\Day2.docx",
    r"C:\Users\王文超\Desktop\Servre_SIT_Mastery\docs\100-day-plan\day-docs\Day5.docx",
]

OUTPUT = r"C:\Users\王文超\Desktop\Servre_SIT_Mastery\docs\100-day-plan\day-docs\analysis.txt"

def classify_paragraph(text):
    text = text.strip()
    if not text:
        return "blank"
    if text.startswith("[root@"):
        return "command"
    # Check for command output indicators (no prompt, often has timestamps, sizes, etc.)
    if any(text.startswith(p) for p in ["#", "$", "drwx", "-rw", "total", "UID", "Last", "Date", "2025-", "2026-"]):
        return "output"
    if text.startswith(("* ", "- ")):
        return "list_item"
    return "normal"

def fmt_size(size_pt):
    if size_pt is None:
        return "None"
    return f"{size_pt}pt"

def get_font_info(run):
    font = run.font
    name = font.name
    # Also check East Asian font
    rFonts = run._element.find(qn('w:rFonts'))
    east_asia = None
    if rFonts is not None:
        east_asia = rFonts.get(qn('w:eastAsia'))

    bold = run.bold
    italic = run.italic

    return {
        "text": repr(run.text),
        "bold": bold,
        "italic": italic,
        "font_name": name,
        "east_asia": east_asia,
        "font_size": fmt_size(font.size),
    }

def analyze_file(filepath, lines):
    filename = os.path.basename(filepath)
    lines.append("=" * 100)
    lines.append(f"FILE: {filename}")
    lines.append(f"PATH: {filepath}")
    lines.append("=" * 100)

    doc = Document(filepath)

    # Check for tables
    tables = doc.tables
    lines.append(f"\nNumber of tables: {len(tables)}")
    for i, table in enumerate(tables):
        lines.append(f"\n  TABLE {i+1}: {len(table.rows)} rows x {len(table.columns)} columns")
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            lines.append(f"    {row_text}")

    # Analyze paragraphs
    lines.append(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    lines.append("")

    # Collect unique styles
    styles_seen = set()

    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "None"
        styles_seen.add(style_name)
        text = para.text.strip()
        classification = classify_paragraph(text)

        # Skip blank paragraphs but note them
        if not text:
            lines.append(f"--- Para {idx+1} [style={style_name}] [blank] ---")
            continue

        lines.append(f"--- Para {idx+1} [style={style_name}] [{classification}] ---")

        # Show first 120 chars of text
        display = text[:150]
        if len(text) > 150:
            display += "..."
        lines.append(f"  TEXT: {display}")

        lines.append(f"  RUNS ({len(para.runs)}):")
        for r_idx, run in enumerate(para.runs):
            info = get_font_info(run)
            if r_idx < 10:  # Only show first 10 runs in detail
                lines.append(f"    Run {r_idx+1}: text={info['text']}, bold={info['bold']}, italic={info['italic']}, font={info['font_name']}, east_asia={info['east_asia']}, size={info['font_size']}")
            elif r_idx == 10:
                lines.append(f"    ... ({len(para.runs) - 10} more runs)")

        lines.append("")

    lines.append(f"Unique styles in this file: {sorted(styles_seen)}")
    lines.append("")

def main():
    lines = []
    for filepath in FILES:
        if os.path.exists(filepath):
            analyze_file(filepath, lines)
        else:
            lines.append(f"FILE NOT FOUND: {filepath}")

    with open(OUTPUT, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    print(f"Analysis written to {OUTPUT}")
    print(f"Total lines: {len(lines)}")

if __name__ == "__main__":
    main()
